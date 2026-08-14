from __future__ import annotations

from pathlib import Path


class ResumeParseError(ValueError):
    pass


def parse_pdf(path: Path) -> str:
    try:
        import fitz
    except Exception as exc:  # pragma: no cover
        raise ResumeParseError("PyMuPDF is not installed.") from exc

    try:
        text = []
        with fitz.open(path) as document:
            for page in document:
                text.append(page.get_text())
        return "\n".join(text).strip()
    except Exception as exc:
        raise ResumeParseError("Unable to parse PDF. It may be corrupted or password protected.") from exc


def parse_docx(path: Path) -> str:
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover
        raise ResumeParseError("python-docx is not installed.") from exc

    try:
        document = Document(path)
        text = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                text.extend(cell.text for cell in row.cells)
        return "\n".join(item for item in text if item).strip()
    except Exception as exc:
        raise ResumeParseError("Unable to parse DOCX. The file may be corrupted.") from exc


def parse_txt(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="ignore").strip()
    except Exception as exc:
        raise ResumeParseError("Unable to read text resume.") from exc


def parse_resume(path: str | Path) -> str:
    file_path = Path(path)
    extension = file_path.suffix.lower()
    if extension == ".pdf":
        text = parse_pdf(file_path)
    elif extension == ".docx":
        text = parse_docx(file_path)
    elif extension == ".txt":
        text = parse_txt(file_path)
    else:
        raise ResumeParseError("Unsupported resume format.")
    if not text:
        raise ResumeParseError("No readable text found in resume.")
    return text

